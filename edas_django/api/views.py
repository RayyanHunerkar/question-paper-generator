import csv
from django.core.files.base import ContentFile
from django.core.files.storage import FileSystemStorage
from django.http import StreamingHttpResponse
from rest_framework.response import Response
from docx import Document
import io
from rest_framework import mixins
from rest_framework import viewsets
from rawquestion.models import RawQuestion
from questiongenerator.models import QuestionPaper, QuestionSet, Question
from .serializers import RawQuestionSerializer, QuestionPaperSerializer, QuestionSetSerializer, QuestionSerializer

fs = FileSystemStorage(location='tmp/')

class RawQuestionViewSet(viewsets.GenericViewSet, 
                mixins.RetrieveModelMixin, 
                mixins.ListModelMixin
    ):
    queryset = RawQuestion.objects.all()
    serializer_class = RawQuestionSerializer

    def create(self, request):
        file = request.FILES["file"]

        csv_content = file.read()
        file_content = ContentFile(csv_content)
        file_name = fs.save(
            "_tmp.csv", file_content
        )
        tmp_file = fs.path(file_name)
        csv_file = open(tmp_file, errors="ignore")
        reader = csv.reader(csv_file)
        next(reader)
        rawquestion_create = []
        for id_,row in enumerate(reader):
            (
                content, 
                difficulty, 
                unit,

            ) = row

            rawquestion_create.append(
                RawQuestion(
                    content=content, 
                    difficulty=difficulty, 
                    unit=unit,
                )
            )

        RawQuestion.objects.bulk_create(rawquestion_create)

        return Response({"status": "success"})

class QuestionPaperViewSet(mixins.ListModelMixin,
                        mixins.CreateModelMixin,
                        mixins.RetrieveModelMixin,
                        mixins.UpdateModelMixin,
                        mixins.DestroyModelMixin,
                        viewsets.GenericViewSet
    ):

    queryset = QuestionPaper.objects.all()
    serializer_class = QuestionPaperSerializer

class QuestionSetViewSet(mixins.ListModelMixin,
                        mixins.CreateModelMixin,
                        viewsets.GenericViewSet):
    queryset = QuestionSet.objects.all()
    serializer_class = QuestionSetSerializer

class ExportDocViewSet(viewsets.GenericViewSet):

    questionsetqueryset = QuestionSet.objects.all()
    questionqueryset = Question.objects.all()
    

    def list(self,request):
        return Response({"status": "success"})

    def retrieve(self, request,pk):
        
        questionpaper = QuestionPaper.objects.all()
        questionset = self.questionsetqueryset.filter(questionpaper=questionpaper)
        question = self.questionqueryset.filter(questionset=questionset)
        document = Document()

        document.add_heading('Question Paper {}'.format(questionpaper.filter(QuestionPaperID=pk).first().QuestionPaperID), 0)
        document.add_paragraph('Academic Year: {}'.format(questionpaper.filter(QuestionPaperID=pk).first().AcademicYear))
        document.add_paragraph('Course: {}'.format(questionpaper.filter(QuestionPaperID=pk).first().subjectcode.name))
        document.add_paragraph('Notes: {}'.format(questionpaper.filter(QuestionPaperID=pk).first().exam_note.note))
        document.add_paragraph('Total Marks: {}'.format(questionpaper.filter(QuestionPaperID=pk).first().maxmarks))
        
        for questionset in self.questionsetqueryset.all().filter(questionpaper_id = pk):
            document.add_paragraph('Question Set: {}'.format(questionset.setType))
            for question in self.questionqueryset.all():
                    if question.questionset == questionset:
                        document.add_paragraph('Question: {}'.format(question.content))
                    


        
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)

        response = StreamingHttpResponse(
            streaming_content=buffer,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

        response['Content-Disposition'] = 'attachment; filename=Question Paper {}.docx'.format('1')
        response['Content-Encoding'] = 'utf-8'
        return response

    '''def list(self, request):
        document = Document()

        document.add_heading('Question Paper {}'.format(self.questionpaperqueryset.first().QuestionPaperID), 0)
        document.add_paragraph('Academic Year: {}'.format(self.questionpaperqueryset.first().AcademicYear.__str__()))
        document.add_paragraph('Course: {}'.format(self.questionpaperqueryset.first().subjectcode.name))
        document.add_paragraph('Notes: {}'.format(self.questionpaperqueryset.first().exam_note.note))
        document.add_paragraph('Total Marks: {}'.format(self.questionpaperqueryset.first().maxmarks))
        
        for questionset in self.questionsetqueryset.all():
            document.add_paragraph('Question Set: {}'.format(questionset.setType))
            for question in self.questionqueryset.all():
                    if question.questionset == questionset:
                        document.add_paragraph('Question: {}'.format(question.content))
                    # document.add_paragraph(question.content)


        
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)

        response = StreamingHttpResponse(
            streaming_content=buffer,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

        response['Content-Disposition'] = 'attachment; filename=Question Paper {}.docx'.format(self.questionpaperqueryset.first().QuestionPaperID)
        response['Content-Encoding'] = 'utf-8'
        return response'''
